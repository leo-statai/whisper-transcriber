from io import BytesIO
from typing import Iterable

from docx import Document
from docx.shared import Pt


def _ts(seconds: float) -> str:
    if seconds is None or seconds < 0:
        seconds = 0
    s = int(seconds)
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def render(job, segments: Iterable) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Transcrição", level=1)
    title.alignment = 1

    meta = doc.add_paragraph()
    meta.add_run("Arquivo: ").bold = True
    meta.add_run(job.filename or "")
    meta.add_run("\nIdioma: ").bold = True
    meta.add_run((job.detected_language or job.language or "—"))
    meta.add_run("\nModelo: ").bold = True
    meta.add_run(job.model or "—")
    if job.duration_seconds:
        meta.add_run("\nDuração: ").bold = True
        meta.add_run(_ts(job.duration_seconds))

    doc.add_paragraph("")
    for seg in segments:
        p = doc.add_paragraph()
        run = p.add_run(f"[{_ts(seg.start)} – {_ts(seg.end)}]  ")
        run.bold = True
        p.add_run(seg.text.strip())

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
